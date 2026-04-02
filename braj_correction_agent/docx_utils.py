"""
docx_utils.py — DOCX read/write helpers for braj correction agent.

Paragraph type is identified by the color of the first run, matching
the colors set by src/generate.py and preserved by src/transliterate_docx.py.
"""

import json
from pathlib import Path
from docx import Document

# Colors that identify Braj commentary paragraphs (hex RGB strings)
BRAJ_COLORS = {"800000", "0000FF"}  # maroon = braj, blue = braj_sub
GURBANI_COLOR = "00007F"            # navy — never touch


def _get_run_color(run) -> str | None:
    """Return the hex RGB color string of a run, or None."""
    try:
        rgb = run.font.color.rgb
        return str(rgb) if rgb else None
    except Exception:
        return None


def extract_braj_paragraphs(docx_path: str) -> list[dict]:
    """
    Read the DOCX and return all Braj/Braj_sub paragraphs.

    Returns a list of dicts:
        {para_idx: int, text: str, color: str}

    para_idx is the paragraph's position in doc.paragraphs (0-based).
    """
    doc = Document(docx_path)
    results = []
    for idx, para in enumerate(doc.paragraphs):
        if not para.runs:
            continue
        color = _get_run_color(para.runs[0])
        if color in BRAJ_COLORS and para.text.strip():
            results.append({"para_idx": idx, "text": para.text, "color": color})
    return results


def build_output_docx(source_path: str, log_path: str, output_path: str) -> int:
    """
    Apply corrections from corrections_log.jsonl to source DOCX and save to output_path.

    Returns the number of corrections applied.
    """
    # Load corrections: last entry per para_idx wins (supports redo)
    corrections: dict[int, str] = {}
    log_file = Path(log_path)
    if log_file.exists():
        with open(log_file) as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                entry = json.loads(line)
                if entry.get("status") == "accepted":
                    corrections[entry["para_idx"]] = entry["corrected"]

    if not corrections:
        print("No accepted corrections found in log. Output DOCX not written.")
        return 0

    doc = Document(source_path)
    applied = 0
    for idx, para in enumerate(doc.paragraphs):
        if idx in corrections:
            # Replace text of the first run, clear others
            if para.runs:
                para.runs[0].text = corrections[idx]
                for run in para.runs[1:]:
                    run.text = ""
                applied += 1

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return applied

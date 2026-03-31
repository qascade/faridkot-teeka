#!/usr/bin/env python3
"""
generate.py — CLI for converting Fareedkot_Teeka.pdf pages to .docx

Usage:
    python3 generate.py --start 338 --end 347
    python3 generate.py -s 338 -e 347
    python3 generate.py -s 338 -e 347 --pdf MyTeeka.pdf
    python3 generate.py -s 338 -e 347 --output output/custom_name.docx
    python3 generate.py -s 338 -e 347 --no-page-markers
    python3 generate.py -s 338 -e 347 --types gurbani braj
    python3 generate.py -s 338 -e 347 --page-break
"""

from __future__ import annotations

import argparse
import logging
import multiprocessing as mp
import os
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from extractor import Element, extract_page, postprocess

# Suppress pdfminer PDF metadata warnings
logging.getLogger('pdfminer').setLevel(logging.CRITICAL)


# ---------------------------------------------------------------------------
# Parallel extraction worker (module-level for multiprocessing pickling)
# ---------------------------------------------------------------------------

def _extract_worker(args: tuple) -> tuple[int, list[Element], str | None]:
    """Extract and postprocess a single page. Returns (page_num, elements, error)."""
    pdf_path, page_num = args
    try:
        elements = postprocess(extract_page(pdf_path, page_num))
        return (page_num, elements, None)
    except Exception as exc:
        return (page_num, [], str(exc))

# ---------------------------------------------------------------------------
# Style config (same as docx_generator.py)
# ---------------------------------------------------------------------------

_GURMUKHI_FONT = 'Gurmukhi MN'
_LATIN_FONT = 'Times New Roman'

_COLORS = {
    'gurbani':    RGBColor(0x00, 0x00, 0x7F),
    'braj':       RGBColor(0x80, 0x00, 0x00),
    'braj_sub':   RGBColor(0x00, 0x00, 0xFF),
    'ggs_ref':    RGBColor(0xC0, 0x00, 0x6A),
    'annotation': RGBColor(0x55, 0x55, 0x55),
    'latin':      RGBColor(0x00, 0x00, 0x00),
    'page_marker': RGBColor(0x99, 0x99, 0x99),
}

_FONT_SIZES = {
    'gurbani':    16,
    'braj':       12,
    'braj_sub':   11,
    'ggs_ref':    11,
    'annotation':  9,
    'latin':      10,
    'page_marker':  8,
}

_ALIGNMENTS = {
    'gurbani':    WD_ALIGN_PARAGRAPH.CENTER,
    'braj':       WD_ALIGN_PARAGRAPH.CENTER,
    'braj_sub':   WD_ALIGN_PARAGRAPH.CENTER,
    'ggs_ref':    WD_ALIGN_PARAGRAPH.LEFT,
    'annotation': WD_ALIGN_PARAGRAPH.LEFT,
    'latin':      WD_ALIGN_PARAGRAPH.LEFT,
    'page_marker': WD_ALIGN_PARAGRAPH.CENTER,
}

ALL_TYPES = {'gurbani', 'braj', 'braj_sub', 'ggs_ref', 'annotation', 'latin'}


def _add_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type())


def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    return br


def _add_paragraph(doc: Document, text: str, text_type: str) -> None:
    color = _COLORS.get(text_type, _COLORS['braj'])
    size = _FONT_SIZES.get(text_type, 12)
    align = _ALIGNMENTS.get(text_type, WD_ALIGN_PARAGRAPH.LEFT)
    font_name = _LATIN_FONT if text_type == 'latin' else _GURMUKHI_FONT

    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(2)
    if text_type == 'gurbani':
        para.paragraph_format.space_before = Pt(4)

    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if text_type == 'gurbani':
        run.font.bold = True

    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)

    if text_type not in ('latin', 'page_marker'):
        lang = OxmlElement('w:lang')
        lang.set(qn('w:val'), 'pa-IN')
        lang.set(qn('w:bidi'), 'pa-IN')
        rPr.append(lang)
        noProof = OxmlElement('w:noProof')
        rPr.append(noProof)


def _add_page_separator(doc: Document, page_num: int) -> None:
    """Add a visible page marker line."""
    _add_paragraph(doc, f'─── Page {page_num} ───', 'page_marker')



def generate(
    pdf_path: str,
    start: int,
    end: int,
    output_path: str,
    page_markers: bool = True,
    page_breaks: bool = False,
    include_types: set[str] | None = None,
    verbose: bool = True,
) -> None:
    """Extract pages start..end (inclusive) and write to output_path using parallel extraction."""
    if include_types is None:
        include_types = ALL_TYPES

    total = end - start + 1

    # Parallel extraction phase
    if verbose:
        print(f"Extracting {total} pages in parallel...", flush=True)

    page_range = list(range(start, end + 1))
    worker_args = [(pdf_path, page_num) for page_num in page_range]

    with mp.Pool() as pool:
        # pool.imap preserves input order AND allows progress feedback
        results = []
        for idx, result in enumerate(pool.imap(_extract_worker, worker_args, chunksize=4), 1):
            if verbose and idx % 50 == 0:
                print(f"  Extracted {idx}/{total} pages")
            results.append(result)

    # Build results dict for easy lookup
    page_data = {}
    for page_num, elements, error in results:
        if error and verbose:
            print(f"Page {page_num}: ERROR: {error}")
        page_data[page_num] = elements

    # Rendering phase (single-threaded, in order)
    if verbose:
        print(f"Rendering {total} pages to docx...")

    # Initialize log file
    log_file = output_path.replace('.docx', '_progress.log')
    if verbose:
        with open(log_file, 'w') as f:
            f.write(f"Rendering pages {start}-{end}\n")

    doc = Document()

    # Remove default empty paragraph
    for para in doc.paragraphs:
        para._element.getparent().remove(para._element)

    # Set narrow margins
    section = doc.sections[0]
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    for i, page_num in enumerate(page_range):
        if verbose:
            print(f"  [{i+1}/{total}] Page {page_num}...", end=' ', flush=True)

        elements = page_data.get(page_num, [])

        if verbose:
            counts = {t: sum(1 for e in elements if e.text_type == t) for t in include_types}
            summary = ' '.join(f'{t}={n}' for t, n in counts.items() if n)
            summary_str = summary or '(empty)'
            print(summary_str)

            # Log to file
            try:
                with open(log_file, 'a') as f:
                    f.write(f"[{i+1}/{total}] Page {page_num}: {summary_str}\n")
            except Exception:
                pass  # Silently skip if can't write log

        if page_markers:
            _add_page_separator(doc, page_num)

        for el in elements:
            if el.text_type not in include_types:
                continue
            # Skip the latin page number element (just digits) — page_marker covers it
            if el.text_type == 'latin' and el.unicode_text.strip().isdigit():
                continue
            text = ' '.join(el.unicode_text.split())
            if not text:
                continue
            _add_paragraph(doc, text, el.text_type)
            # Blank line after every braj block
            if el.text_type in ('braj',):
                doc.add_paragraph()

        if page_breaks and i < total - 1:
            para = doc.add_paragraph()
            run = para.add_run()
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            run._r.append(br)

    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    print(f"\nWritten: {output_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        description='Convert Fareedkot_Teeka.pdf pages to Unicode .docx',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python3 generate.py -s 338 -e 347
  python3 generate.py -s 338 -e 347 --page-break
  python3 generate.py -s 1 -e 50 --output output/first_50.docx
  python3 generate.py -s 338 -e 347 --types gurbani braj
  python3 generate.py -s 338 -e 347 --no-page-markers
        """,
    )
    p.add_argument('-s', '--start', type=int, required=True, help='Start page (1-based)')
    p.add_argument('-e', '--end', type=int, required=True, help='End page (inclusive)')
    p.add_argument('--pdf', default='Fareedkot_Teeka.pdf', help='Input PDF path (default: Fareedkot_Teeka.pdf)')
    p.add_argument('--output', default=None, help='Output .docx path (default: output/pages_START-END.docx)')
    p.add_argument('--no-page-markers', action='store_true', help='Omit the ─── Page N ─── separator lines')
    p.add_argument('--page-break', action='store_true', help='Insert a page break between pages')
    p.add_argument(
        '--types', nargs='+',
        choices=sorted(ALL_TYPES),
        default=None,
        help='Only include these text types (default: all)',
    )
    p.add_argument('-q', '--quiet', action='store_true', help='Suppress per-page output')
    return p


def main() -> None:
    parser = build_parser()
    args = parser.parse_args()

    if args.end < args.start:
        parser.error(f"--end ({args.end}) must be >= --start ({args.start})")

    output = args.output or os.path.join('output', f'pages_{args.start}-{args.end}.docx')
    include_types = set(args.types) if args.types else None

    print(f"Fareedkot Teeka converter")
    print(f"  Pages  : {args.start}–{args.end}  ({args.end - args.start + 1} pages)")
    print(f"  PDF    : {args.pdf}")
    print(f"  Output : {output}")
    if include_types:
        print(f"  Types  : {', '.join(sorted(include_types))}")
    print()

    generate(
        pdf_path=args.pdf,
        start=args.start,
        end=args.end,
        output_path=output,
        page_markers=not args.no_page_markers,
        page_breaks=args.page_break,
        include_types=include_types,
        verbose=not args.quiet,
    )


if __name__ == '__main__':
    main()

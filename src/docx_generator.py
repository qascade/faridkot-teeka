"""
docx_generator.py — Produce a formatted .docx from extracted Element objects.

Color scheme mirrors the original Fareedkot_Teeka.pdf:
  gurbani    → deep blue  (#00007F)
  braj       → deep red   (#800000)
  braj_sub   → bright blue (#0000FF)
  ggs_ref    → deep pink  (#C0006A)
  annotation → gray       (#555555)
  latin      → black
"""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

if TYPE_CHECKING:
    from extractor import Element

_GURMUKHI_FONT = 'Gurmukhi MN'
_LATIN_FONT = 'Times New Roman'

_COLORS = {
    'gurbani':    RGBColor(0x00, 0x00, 0x7F),
    'braj':       RGBColor(0x80, 0x00, 0x00),
    'braj_sub':   RGBColor(0x00, 0x00, 0xFF),
    'ggs_ref':    RGBColor(0xC0, 0x00, 0x6A),
    'annotation': RGBColor(0x55, 0x55, 0x55),
    'latin':      RGBColor(0x00, 0x00, 0x00),
}

_FONT_SIZES = {
    'gurbani':    16,
    'braj':       12,
    'braj_sub':   11,
    'ggs_ref':    11,
    'annotation':  9,
    'latin':      10,
}

_ALIGNMENTS = {
    'gurbani':    WD_ALIGN_PARAGRAPH.CENTER,
    'braj':       WD_ALIGN_PARAGRAPH.CENTER,
    'braj_sub':   WD_ALIGN_PARAGRAPH.CENTER,
    'ggs_ref':    WD_ALIGN_PARAGRAPH.LEFT,
    'annotation': WD_ALIGN_PARAGRAPH.LEFT,
    'latin':      WD_ALIGN_PARAGRAPH.LEFT,
}


def _add_paragraph(doc: Document, text: str, text_type: str) -> None:
    color = _COLORS.get(text_type, _COLORS['braj'])
    size = _FONT_SIZES.get(text_type, 12)
    align = _ALIGNMENTS.get(text_type, WD_ALIGN_PARAGRAPH.LEFT)
    font_name = _LATIN_FONT if text_type == 'latin' else _GURMUKHI_FONT

    para = doc.add_paragraph()
    para.alignment = align

    # Remove space after paragraph
    para.paragraph_format.space_after = Pt(2)
    if text_type == 'gurbani':
        para.paragraph_format.space_before = Pt(6)

    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if text_type == 'gurbani':
        run.font.bold = True

    # Ensure the font applies to complex scripts (Gurmukhi)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)

    if text_type != 'latin':
        # Language tag: tells Word to apply Gurmukhi shaping rules (fixes sihari/matra ordering visually)
        lang = OxmlElement('w:lang')
        lang.set(qn('w:val'), 'pa-IN')
        lang.set(qn('w:bidi'), 'pa-IN')
        rPr.append(lang)
        # Disable spell-check for Gurmukhi runs (no Punjabi dictionary in most installs)
        noProof = OxmlElement('w:noProof')
        rPr.append(noProof)


def create_docx(elements: list[Element], output_path: str) -> None:
    """Render *elements* into a new .docx at *output_path*."""
    doc = Document()

    # Set default margins (narrow)
    section = doc.sections[0]
    from docx.shared import Mm
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    # Remove the default empty paragraph
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    for el in elements:
        text = ' '.join(el.unicode_text.split())
        if not text:
            continue
        _add_paragraph(doc, text, el.text_type)

    doc.save(output_path)
    print(f"Written: {output_path}")
